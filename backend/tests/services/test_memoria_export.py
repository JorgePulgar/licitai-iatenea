"""Tests del servicio de export de la Memoria (5.6 ♻): HTML, PDF (WeasyPrint) y DOCX.

Los tests de PDF ejercitan WeasyPrint real (libs nativas presentes vía dev.sh); si
faltasen, se saltan — el endpoint ya degrada a 503 en ese caso.
"""

import io

import pytest

from app.services.memoria_export import (
    ExportOptions,
    normalize_markdown,
    render_markdown_docx,
    render_markdown_html,
)

try:
    from weasyprint import HTML as _WeasyHTML  # noqa: N814

    _WeasyHTML(string="<p>x</p>").write_pdf()
    HAS_WEASYPRINT = True
except Exception:
    HAS_WEASYPRINT = False

requires_weasyprint = pytest.mark.skipif(
    not HAS_WEASYPRINT, reason="WeasyPrint sin libs nativas (Pango) en esta máquina"
)

# PNG 1×1 rojo, válido — para logo e imágenes embebidas.
_PNG_DATA_URI = (
    "data:image/png;base64,"
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR4nGP4z8DwHwAFAAH/q842"
    "iQAAAABJRU5ErkJggg=="
)


# ── Normalización ────────────────────────────────────────────────────────────


def test_normalize_markdown_removes_document_fence():
    markdown = "```markdown\n# Memoria\n\nTexto con **formato**.\n```"

    assert normalize_markdown(markdown) == "# Memoria\n\nTexto con **formato**."


def test_normalize_markdown_preserves_regular_code_blocks():
    markdown = "# Memoria\n\n```python\nprint('ok')\n```"

    assert normalize_markdown(markdown) == markdown


# ── HTML de impresión (contrato con el editor del frontend) ──────────────────


def test_render_markdown_html_processes_wrapped_markdown():
    html = render_markdown_html("```md\n# Memoria\n\nTexto con **formato**.\n```")

    assert "<h1" in html and "Memoria</h1>" in html
    assert "<strong>formato</strong>" in html
    assert "language-md" not in html
    assert 'content: "Página " counter(page) " de " counter(pages)' in html


def test_render_markdown_html_keeps_document_components():
    markdown = """
<header class="document-header">Cabecera corporativa</header>

# Memoria

<figure class="document-video">
  <span class="document-video__label">Vídeo de presentación</span>
  <a href="https://example.com/video">https://example.com/video</a>
</figure>

<footer class="document-footer">Oferta técnica</footer>
"""

    html = render_markdown_html(markdown)

    assert '<header class="document-header">Cabecera corporativa</header>' in html
    assert '<figure class="document-video">' in html
    assert '<footer class="document-footer">Oferta técnica</footer>' in html
    assert html.index('<header class="document-header">') < html.index("Memoria</h1>")
    assert html.index('<footer class="document-footer">') < html.index("Memoria</h1>")


def test_render_markdown_html_resolves_variables_and_page_settings():
    markdown = """
<div class="document-settings" data-page-size="letter" data-orientation="landscape"></div>
<header class="document-header">
  <span class="document-variable" data-variable="company_name">Nombre de empresa</span>
</header>
<footer class="document-footer">
  Página <span class="document-variable" data-variable="page">Página actual</span>
  de <span class="document-variable" data-variable="pages">Total de páginas</span>
</footer>

# Memoria
"""

    html = render_markdown_html(markdown, {"company_name": "Empresa & Asociados"})

    assert "@page { size: Letter landscape;" in html
    assert "Empresa &amp; Asociados" in html
    assert 'data-variable="page"></span>' in html
    assert 'data-variable="pages"></span>' in html
    assert 'content: "Página " counter(page) " de " counter(pages)' not in html


def test_render_markdown_html_keeps_image_layout_attributes():
    markdown = """
<figure class="document-image" data-width="50" data-rotation="90" data-align="right"
 style="width:50%;margin-left:auto;margin-right:0;">
  <img src="data:image/png;base64,abc" alt="Plano" style="transform:rotate(90deg);">
</figure>
"""

    html = render_markdown_html(markdown)

    assert 'data-width="50"' in html
    assert "transform:rotate(90deg)" in html


# ── Opciones 5.6: header/footer por defecto y TOC ────────────────────────────


def test_default_header_footer_injected_from_options():
    html = render_markdown_html(
        "# Memoria\n\nTexto.",
        options=ExportOptions(
            header_text="ACME S.L.", footer_text="Oferta técnica",
            logo_data_uri=_PNG_DATA_URI,
        ),
    )

    assert '<header class="document-header">' in html
    assert "ACME S.L." in html
    assert '<img class="document-logo"' in html
    assert '<footer class="document-footer">Oferta técnica</footer>' in html
    # Los running van antes del contenido.
    assert html.index('<header class="document-header">') < html.index("Memoria</h1>")


def test_document_components_win_over_options():
    markdown = '<header class="document-header">La del documento</header>\n\n# Memoria'
    html = render_markdown_html(
        markdown, options=ExportOptions(header_text="La de las opciones")
    )

    assert "La del documento" in html
    assert "La de las opciones" not in html


def test_toc_built_from_h2_h3_with_page_counters():
    markdown = (
        "# Memoria\n\n## Introducción\n\nTexto.\n\n### Contexto\n\nTexto.\n\n"
        "## Solución propuesta\n\nTexto."
    )
    html = render_markdown_html(markdown, options=ExportOptions(include_toc=True))

    assert '<nav class="document-toc">' in html
    assert "Índice" in html
    assert 'class="toc-l2"><a href="#introduccion">Introducción</a>' in html
    assert 'class="toc-l3"><a href="#contexto">Contexto</a>' in html
    assert "target-counter(attr(href), page)" in html  # nº de página real en PDF
    # El índice va antes del contenido y tras los running.
    assert html.index('<nav class="document-toc">') < html.index("Introducción</h2>")


def test_toc_omitted_without_flag_and_without_headings():
    assert '<nav class="document-toc">' not in render_markdown_html("# Solo título")
    assert '<nav class="document-toc">' not in render_markdown_html(
        "# Solo título", options=ExportOptions(include_toc=True)
    )


# ── PDF (aceptación 5.6: multipágina con tablas; rótulos nunca huérfanos) ────


def _long_memoria_with_tables(sections: int = 8, rows: int = 12) -> str:
    parts = ["# Memoria técnica\n"]
    for n in range(1, sections + 1):
        parts.append(f"## Apartado {n}\n")
        parts.append(("Párrafo de contexto del apartado. " * 12 + "\n\n") * 2)
        parts.append(f"Tabla {n}: desglose de medios del apartado {n}.\n")
        parts.append("| Perfil | Dedicación | Certificación |\n|---|---|---|\n")
        for r in range(1, rows + 1):
            parts.append(f"| Perfil {n}.{r} | {r * 10}% | CERT-{n}{r:02d} |\n")
        parts.append("\n")
    return "".join(parts)


@requires_weasyprint
def test_pdf_multipage_with_tables_and_captions_not_orphaned():
    from pypdf import PdfReader

    pdf = io.BytesIO(
        __import__("app.services.memoria_export", fromlist=["render_markdown_pdf"])
        .render_markdown_pdf(_long_memoria_with_tables())
    )
    reader = PdfReader(pdf)
    pages_text = [page.extract_text() or "" for page in reader.pages]

    assert len(pages_text) >= 2, "la memoria de aceptación debe ocupar varias páginas"
    # Rótulo nunca huérfano: cada "Tabla N:" convive en su página con la primera
    # fila de su tabla (Perfil N.1).
    for n in range(1, 9):
        caption_pages = [i for i, t in enumerate(pages_text) if f"Tabla {n}:" in t]
        assert caption_pages, f"rótulo de la tabla {n} no encontrado"
        assert any(
            f"Perfil {n}.1" in pages_text[i] for i in caption_pages
        ), f"rótulo 'Tabla {n}:' huérfano — su tabla empieza en otra página"


@requires_weasyprint
def test_pdf_with_toc_renders():
    from app.services.memoria_export import render_markdown_pdf

    pdf = render_markdown_pdf(
        _long_memoria_with_tables(sections=3),
        options=ExportOptions(include_toc=True, header_text="ACME S.L."),
    )
    assert pdf[:5] == b"%PDF-"


# ── DOCX (nuevo en 5.6) ──────────────────────────────────────────────────────


def _load_docx(data: bytes):
    from docx import Document

    return Document(io.BytesIO(data))


def test_docx_headings_paragraphs_and_inline_formatting():
    markdown = (
        "# Memoria técnica\n\n## Introducción\n\n"
        "Texto con **negrita**, *cursiva* y `codigo`.\n"
    )
    doc = _load_docx(render_markdown_docx(markdown))
    styles = [p.style.name for p in doc.paragraphs]
    texts = [p.text for p in doc.paragraphs]

    assert "Heading 1" in styles and "Heading 2" in styles
    assert any("Texto con negrita, cursiva y codigo." == t for t in texts)
    body = next(p for p in doc.paragraphs if p.text.startswith("Texto con"))
    formats = {r.text: (r.bold, r.italic, r.font.name) for r in body.runs}
    assert formats["negrita"][0] is True
    assert formats["cursiva"][1] is True
    assert formats["codigo"][2] == "Courier New"


def test_docx_lists_nested():
    markdown = "- uno\n- dos\n    - dos punto uno\n\n1. primero\n2. segundo\n"
    doc = _load_docx(render_markdown_docx(markdown))
    styles = [(p.style.name, p.text) for p in doc.paragraphs if p.text]

    assert ("List Bullet", "uno") in styles
    assert ("List Bullet 2", "dos punto uno") in styles
    assert ("List Number", "primero") in styles


def test_docx_table_with_header_row():
    markdown = "| Perfil | Años |\n|---|---|\n| Jefe de proyecto | 10 |\n"
    doc = _load_docx(render_markdown_docx(markdown))

    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.cell(0, 0).text == "Perfil"
    assert table.cell(1, 0).text == "Jefe de proyecto"
    assert table.cell(1, 1).text == "10"
    header_run = table.cell(0, 0).paragraphs[0].runs[0]
    assert header_run.bold is True


def test_docx_code_block_monospace():
    # Con encabezado delante: un fence que ocupa TODO el documento se trata como
    # envoltorio global del modelo y se retira (normalize_markdown).
    doc = _load_docx(render_markdown_docx("# Memoria\n\n```\ncomando --flag\n```\n"))
    code = next(p for p in doc.paragraphs if "comando --flag" in p.text)

    assert code.runs[0].font.name == "Courier New"


def test_docx_header_footer_and_page_fields():
    doc = _load_docx(
        render_markdown_docx(
            "# Memoria",
            options=ExportOptions(
                header_text="ACME S.L.", footer_text="Oferta técnica",
                logo_data_uri=_PNG_DATA_URI,
            ),
        )
    )
    section = doc.sections[0]

    assert "ACME S.L." in section.header.paragraphs[0].text
    footer_xml = section.footer.paragraphs[0]._p.xml
    assert "Oferta técnica" in section.footer.paragraphs[0].text
    assert "PAGE" in footer_xml and "NUMPAGES" in footer_xml
    # Logo embebido en la cabecera.
    assert section.header.part.related_parts, "el logo debe quedar embebido"


def test_docx_document_components_feed_header_and_are_not_body():
    markdown = (
        '<header class="document-header">Cabecera del documento</header>\n\n'
        "# Memoria\n\nTexto.\n\n"
        '<footer class="document-footer">Pie del documento</footer>\n'
    )
    doc = _load_docx(render_markdown_docx(markdown))

    assert "Cabecera del documento" in doc.sections[0].header.paragraphs[0].text
    assert "Pie del documento" in doc.sections[0].footer.paragraphs[0].text
    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Cabecera del documento" not in body_text
    assert "Pie del documento" not in body_text


def test_docx_toc_static_list():
    markdown = "# Memoria\n\n## Introducción\n\nTexto.\n\n### Contexto\n\nTexto.\n"
    doc = _load_docx(render_markdown_docx(markdown, options=ExportOptions(include_toc=True)))
    texts = [p.text for p in doc.paragraphs]

    índice = texts.index("Índice")
    assert "Introducción" in texts[índice + 1 :]
    assert "Contexto" in texts[índice + 1 :]
    # Y los encabezados reales del cuerpo siguen presentes después.
    assert sum(1 for t in texts if t == "Introducción") == 2  # TOC + encabezado


def test_docx_embedded_image():
    markdown = f"# Memoria\n\n![Plano]({_PNG_DATA_URI})\n"
    doc = _load_docx(render_markdown_docx(markdown))

    assert doc.inline_shapes, "la imagen data-URI debe quedar embebida"


def test_docx_resolves_document_variables():
    markdown = (
        "# Memoria\n\nPreparado por "
        '<span class="document-variable" data-variable="company_name">X</span>.\n'
    )
    doc = _load_docx(render_markdown_docx(markdown, variables={"company_name": "ACME"}))

    assert any("Preparado por ACME." in p.text for p in doc.paragraphs)
