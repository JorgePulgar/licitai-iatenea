"""Export de la Memoria Técnica: Markdown → HTML de impresión → PDF (WeasyPrint) y DOCX (python-docx).

Reescrito desde spec funcional (tarea 5.6 ♻). Spec:

ENTRADAS
- Markdown de la memoria. Puede venir envuelto en un fence global ```markdown
  (hábito de algunos modelos): se retira antes de renderizar.
- El Markdown puede contener componentes HTML del editor, que son CONTRATO con el
  frontend (misma nomenclatura que `frontend/src/lib/memoriaPaged.css`):
    · `<header class="document-header">` / `<footer class="document-footer">` —
      cabecera/pie repetidos en cada página (elementos running de CSS paged media).
    · `<div class="document-settings" data-page-size="a4|letter|a3"
       data-orientation="portrait|landscape">` — configuración de página.
    · `<span class="document-variable" data-variable="clave">` — variables de
      documento; `page`/`pages` se resuelven a contadores de página, el resto con
      el mapping `variables` (escapado HTML).
    · `<figure class="document-image">` (con `data-*`/`style` de layout) y
      `<figure class="document-video">` — se conservan tal cual.
- `variables`: mapping clave→valor para los spans de variables.
- `options` (`ExportOptions`, todo opcional — NUEVO en 5.6):
    · `header_text` / `footer_text` / `logo_data_uri`: cabecera/pie por defecto
      (nombre de empresa, logo). Solo se inyectan si el documento NO trae su
      propio componente document-header/footer — lo editado en el documento manda.
    · `include_toc`: índice generado desde los encabezados h2/h3. En PDF, con
      número de página real (target-counter) y salto de página tras el índice.
      En DOCX, lista estática indentada (sin números de página: un campo TOC de
      Word exigiría "actualizar campos" al abrir).

SALIDAS
- `render_markdown_html`  → documento HTML completo con CSS de impresión A4.
- `render_markdown_pdf`   → bytes PDF (WeasyPrint).
- `render_markdown_docx`  → bytes DOCX (python-docx): encabezados, párrafos con
  formato inline (negrita/cursiva/código), listas anidadas, tablas con cabecera,
  bloques de código, citas, imágenes data-URI y saltos de página. Cabecera/pie de
  sección desde `options` (o del texto de los componentes del documento) y pie
  con "Página X de Y" mediante campos PAGE/NUMPAGES nativos.

REGLAS DE PAGINACIÓN (aprendizajes 2026-06-23 elevados a requisitos de spec)
- `break-after: avoid` en h1-h3: un encabezado nunca queda huérfano a final de página.
- `break-inside: avoid` en tabla/figura/imagen/pre/blockquote.
- El rótulo ("Tabla 1: …", "Figura 2: …") inmediatamente anterior a una tabla o
  figura no debe separarse de ella (`p:has(+ table) { break-after: avoid }`).
- Contador "Página X de Y" abajo-derecha por defecto; se omite si el documento ya
  usa las variables `page`/`pages` en su propio pie.

Los imports de `markdown`, `weasyprint` y `docx` son perezosos para que la app
cargue aunque falten las librerías nativas de WeasyPrint (Pango). Ver ADR-002 §6.

IMPORTANTE: el CSS debe mantenerse en sintonía con `frontend/src/lib/memoriaPaged.css`
(el visor de la pestaña Memoria aplica las mismas reglas vía paged.js).
"""

import base64
import html
import io
import re
from dataclasses import dataclass
from html.parser import HTMLParser
from typing import Any, Mapping

# ── Opciones de export (5.6) ─────────────────────────────────────────────────


@dataclass(frozen=True)
class ExportOptions:
    header_text: str | None = None
    footer_text: str | None = None
    logo_data_uri: str | None = None
    include_toc: bool = False


# ── Normalización del Markdown ───────────────────────────────────────────────

_GLOBAL_FENCE = re.compile(
    r"\A\s*```(?:markdown|md)?[ \t]*\r?\n(?P<inner>[\s\S]*?)\r?\n```[ \t]*\s*\Z",
    re.IGNORECASE,
)


def normalize_markdown(markdown_text: str) -> str:
    """Retira el fence global ```markdown que algunos modelos añaden al documento."""
    wrapped = _GLOBAL_FENCE.match(markdown_text)
    return wrapped.group("inner") if wrapped else markdown_text


# ── Componentes del editor (contrato con el frontend) ────────────────────────

_RUNNING_COMPONENT = re.compile(
    r"<(?P<tag>header|footer)\b"
    r"(?=[^>]*\bclass=[\"'][^\"']*\bdocument-(?:header|footer)\b[^\"']*[\"'])"
    r"[^>]*>[\s\S]*?</(?P=tag)>",
    re.IGNORECASE,
)
_SETTINGS_COMPONENT = re.compile(
    r"<div\b(?=[^>]*\bclass=[\"'][^\"']*\bdocument-settings\b[^\"']*[\"'])"
    r"(?P<attrs>[^>]*)></div>",
    re.IGNORECASE,
)
_VARIABLE_COMPONENT = re.compile(
    r"<span\b(?P<attrs>[^>]*\bdata-variable=[\"'](?P<key>[a-z_]+)[\"'][^>]*)>"
    r"[\s\S]*?</span>",
    re.IGNORECASE,
)
_HTML_ATTR = re.compile(r"\b(?P<name>[a-z-]+)=[\"'](?P<value>[^\"']*)[\"']", re.IGNORECASE)
_PAGE_VARIABLE = re.compile(r"\bdata-variable=[\"'](?:page|pages)[\"']", re.IGNORECASE)

_PAGE_SIZES = {"a4": "A4", "letter": "Letter", "a3": "A3"}


def _settings_attrs(html_body: str) -> dict[str, str]:
    found = _SETTINGS_COMPONENT.search(html_body)
    if not found:
        return {}
    return {
        attr.group("name").lower(): attr.group("value").lower()
        for attr in _HTML_ATTR.finditer(found.group("attrs"))
    }


def _resolve_variables(html_body: str, variables: Mapping[str, str] | None) -> str:
    """Sustituye las variables de documento; `page`/`pages` quedan vacías (contador CSS)."""
    values = variables or {}

    def substitute(found: re.Match[str]) -> str:
        key = found.group("key").lower()
        attrs = found.group("attrs")
        if key in ("page", "pages"):
            return f"<span{attrs}></span>"
        return f"<span{attrs}>{html.escape(str(values.get(key, '')))}</span>"

    return _VARIABLE_COMPONENT.sub(substitute, html_body)


def _hoist_running_components(html_body: str) -> str:
    """Coloca header/footer running al principio, para que apliquen desde la página 1."""
    components = [c.group(0) for c in _RUNNING_COMPONENT.finditer(html_body)]
    if not components:
        return html_body
    return "".join(components) + _RUNNING_COMPONENT.sub("", html_body).lstrip()


def _default_running_components(html_body: str, options: ExportOptions) -> str:
    """Inyecta cabecera/pie desde options SOLO si el documento no trae los suyos."""
    has_header = bool(re.search(r'\bclass=["\'][^"\']*\bdocument-header\b', html_body))
    has_footer = bool(re.search(r'\bclass=["\'][^"\']*\bdocument-footer\b', html_body))
    prefix = ""
    if not has_header and (options.header_text or options.logo_data_uri):
        logo = (
            f'<img class="document-logo" src="{html.escape(options.logo_data_uri, quote=True)}" alt="">'
            if options.logo_data_uri
            else ""
        )
        text = html.escape(options.header_text or "")
        prefix += f'<header class="document-header">{logo}<span>{text}</span></header>'
    if not has_footer and options.footer_text:
        prefix += (
            f'<footer class="document-footer">{html.escape(options.footer_text)}</footer>'
        )
    return prefix + html_body if prefix else html_body


# ── Índice (TOC) ─────────────────────────────────────────────────────────────


def _walk_toc(toc_tokens: list[dict[str, Any]]):
    """Recorre en orden de documento el árbol de tokens de la extensión `toc`."""
    for token in toc_tokens:
        yield token
        yield from _walk_toc(token.get("children", []))


def _toc_entries(toc_tokens: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Solo h2/h3 (el h1 es el título del documento, no entra en el índice)."""
    return [t for t in _walk_toc(toc_tokens) if t["level"] in (2, 3)]


def _build_toc_nav(toc_tokens: list[dict[str, Any]]) -> str:
    """Índice desde los encabezados h2/h3.

    Cada entrada enlaza a su encabezado; el CSS de impresión resuelve el número
    de página real con target-counter.
    """
    entries = [
        f'<li class="toc-l{t["level"]}">'
        f'<a href="#{t["id"]}">{html.escape(t["name"])}</a></li>'
        for t in _toc_entries(toc_tokens)
    ]
    if not entries:
        return ""
    return (
        '<nav class="document-toc"><h2 class="toc-title">Índice</h2>'
        f"<ol>{''.join(entries)}</ol></nav>"
    )


def _flatten_toc(toc_tokens: list[dict[str, Any]]) -> list[tuple[int, str]]:
    """[(nivel, texto)] de los encabezados h2/h3 — para el índice estático del DOCX."""
    return [(t["level"], t["name"]) for t in _toc_entries(toc_tokens)]


# ── CSS de impresión ─────────────────────────────────────────────────────────
# Requisitos de paginación del spec (ver docstring). Nomenclatura de clases =
# contrato con los componentes del editor y con memoriaPaged.css.

_PRINT_CSS = """
@page {
  size: A4;
  margin: 2.5cm 2cm 2.4cm;
  @top-center { content: element(document-header); }
  @bottom-left { content: element(document-footer); }
}
* { box-sizing: border-box; }
html, body { width: 100%; max-width: 100%; }
body {
  font-family: 'Helvetica', 'Arial', sans-serif;
  font-size: 11pt;
  line-height: 1.5;
  color: #1a1a1a;
  overflow-wrap: anywhere;
  word-break: normal;
}

/* Cabecera y pie repetidos (elementos running) */
.document-header {
  position: running(document-header);
  width: 100%;
  color: #4b5563;
  border-bottom: 0.5pt solid #d1d5db;
  padding-bottom: 5pt;
  font-size: 9pt;
}
.document-header .document-logo {
  display: inline-block;
  max-height: 28pt;
  width: auto;
  margin: 0 8pt 0 0;
  vertical-align: middle;
}
.document-footer {
  position: running(document-footer);
  width: 100%;
  color: #6b7280;
  font-size: 8.5pt;
}

/* Tipografía */
h1 { font-size: 20pt; margin: 0 0 0.6em; }
h2 { font-size: 15pt; margin: 1.2em 0 0.4em; border-bottom: 1px solid #ddd; padding-bottom: 0.2em; }
h3 { font-size: 12.5pt; margin: 1em 0 0.3em; }
p, li { margin: 0.3em 0; }
ul, ol { padding-left: 1.5em; }
a { color: #1d4ed8; overflow-wrap: anywhere; }
blockquote {
  margin: 0.8em 0;
  padding: 0.2em 0 0.2em 1em;
  border-left: 3px solid #d1d5db;
  color: #4b5563;
}
code {
  background: #f4f4f4;
  padding: 1px 4px;
  border-radius: 3px;
  overflow-wrap: anywhere;
}
pre {
  max-width: 100%;
  padding: 8pt;
  background: #f4f4f4;
  white-space: pre-wrap;
  overflow-wrap: anywhere;
  word-break: break-word;
}
pre code { padding: 0; }

/* Imágenes y figuras del editor */
img, svg, video {
  display: block;
  max-width: 100%;
  height: auto;
  object-fit: contain;
  margin: 0.8em auto;
}
figure.document-image {
  display: block;
  max-width: 100%;
  margin-top: 0.8em;
  margin-bottom: 0.8em;
}
figure.document-image img {
  width: 100%;
  max-width: 100%;
  margin: 0;
  transform-origin: center;
}
.document-video {
  max-width: 100%;
  margin: 0.9em 0;
  padding: 12pt;
  border: 1px solid #d1d5db;
  border-radius: 5pt;
  background: #f9fafb;
}
.document-video__label {
  display: block;
  margin-bottom: 3pt;
  color: #111827;
  font-weight: bold;
}
.document-video a {
  display: block;
  font-size: 9pt;
  overflow-wrap: anywhere;
}

/* Tablas */
table {
  border-collapse: collapse;
  width: 100%;
  max-width: 100%;
  table-layout: fixed;
  margin: 0.6em 0;
}
th, td {
  border: 1px solid #ccc;
  padding: 4px 8px;
  text-align: left;
  font-size: 10pt;
  overflow-wrap: anywhere;
  word-break: break-word;
}
th { background: #f2f2f2; }

/* Variables de documento: page/pages via contadores CSS */
.document-variable[data-variable="page"],
.document-variable[data-variable="pages"] {
  font-size: 0;
}
.document-variable[data-variable="page"]::after,
.document-variable[data-variable="pages"]::after {
  font-size: 8.5pt;
}
.document-variable[data-variable="page"]::after { content: counter(page); }
.document-variable[data-variable="pages"]::after { content: counter(pages); }
.document-settings { display: none; }

/* Índice (TOC) con número de página real */
.document-toc { break-after: page; page-break-after: always; }
.document-toc .toc-title { border-bottom: none; }
.document-toc ol { list-style: none; padding-left: 0; }
.document-toc li.toc-l3 { padding-left: 1.5em; }
.document-toc a { color: #1a1a1a; text-decoration: none; }
.document-toc a::after {
  content: leader('.') target-counter(attr(href), page);
  color: #6b7280;
}

/* Reglas de paginación (requisitos del spec — ver docstring del módulo) */
.page-break { break-after: page; page-break-after: always; }
h1, h2, h3 { break-after: avoid; page-break-after: avoid; }
img, figure, table, pre, blockquote {
  break-inside: avoid;
  page-break-inside: avoid;
}
/* El rótulo ("Tabla 1: …", "Figura 2: …") no se separa del elemento que describe */
p:has(+ table), p:has(+ figure),
ul:has(+ table), ol:has(+ table),
li:has(+ table) {
  break-after: avoid;
  page-break-after: avoid;
}
"""


def _page_rule(html_body: str) -> str:
    """Regla @page según document-settings + contador por defecto si procede."""
    attrs = _settings_attrs(html_body)
    size = _PAGE_SIZES.get(attrs.get("data-page-size", "a4"), "A4")
    orientation = "landscape" if attrs.get("data-orientation") == "landscape" else "portrait"
    counter = (
        ""
        if _PAGE_VARIABLE.search(html_body)
        else """
  @bottom-right {
    content: "Página " counter(page) " de " counter(pages);
    color: #6b7280;
    font-size: 8.5pt;
  }"""
    )
    return f"@page {{ size: {size} {orientation};{counter}\n}}"


# ── Render HTML / PDF ────────────────────────────────────────────────────────


def _markdown_to_body(markdown_text: str) -> tuple[str, list[dict[str, Any]]]:
    """Markdown → (HTML del cuerpo, tokens de TOC). La extensión `toc` asigna ids."""
    import markdown as md

    converter = md.Markdown(extensions=["tables", "fenced_code", "sane_lists", "toc"])
    body = converter.convert(normalize_markdown(markdown_text))
    return body, getattr(converter, "toc_tokens", [])


def render_markdown_html(
    markdown_text: str,
    variables: Mapping[str, str] | None = None,
    options: ExportOptions | None = None,
) -> str:
    """Convierte el Markdown de la memoria en un documento HTML de impresión completo."""
    opts = options or ExportOptions()
    body, toc_tokens = _markdown_to_body(markdown_text)
    page_rule = _page_rule(body)
    body = _resolve_variables(body, variables)
    body = _default_running_components(body, opts)
    body = _hoist_running_components(body)
    if opts.include_toc:
        toc_nav = _build_toc_nav(toc_tokens)
        if toc_nav:
            body = _insert_after_running(body, toc_nav)
    return (
        "<!DOCTYPE html><html><head><meta charset='utf-8'>"
        f"<style>{_PRINT_CSS}\n{page_rule}</style></head><body>{body}</body></html>"
    )


def _insert_after_running(body: str, fragment: str) -> str:
    """Inserta un fragmento tras los componentes running ya izados al principio."""
    offset = 0
    for component in _RUNNING_COMPONENT.finditer(body):
        if component.start() != offset:
            break
        offset = component.end()
    return body[:offset] + fragment + body[offset:]


def render_markdown_pdf(
    markdown_text: str,
    variables: Mapping[str, str] | None = None,
    options: ExportOptions | None = None,
) -> bytes:
    """Markdown → PDF (WeasyPrint)."""
    from weasyprint import HTML

    return HTML(string=render_markdown_html(markdown_text, variables, options)).write_pdf()


# ── Render DOCX ──────────────────────────────────────────────────────────────

_DATA_URI = re.compile(r"\Adata:image/[a-z+.-]+;base64,(?P<b64>[A-Za-z0-9+/=\s]+)\Z")
_TAG_STRIP = re.compile(r"<[^>]+>")


@dataclass
class _Run:
    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False


class _BlockCollector(HTMLParser):
    """Recorre el HTML del cuerpo y lo reduce a bloques neutrales para el DOCX.

    Bloques: ("heading", level, runs) · ("para", runs) · ("item", ordered, depth, runs)
    · ("table", rows) · ("codeblock", text) · ("quote", runs) · ("image", bytes)
    · ("pagebreak",). Los componentes del editor (header/footer/settings/toc) se
    omiten: cabecera y pie van a la sección del DOCX, no al cuerpo.
    """

    _SKIPPED = ("document-header", "document-footer", "document-settings", "document-toc",
                "document-video")
    # Elementos void HTML: llegan sin endtag, no deben tocar el contador de skip.
    _VOID = ("br", "img", "hr", "input", "meta", "source", "wbr")

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[tuple[Any, ...]] = []
        self._runs: list[_Run] | None = None
        self._bold = 0
        self._italic = 0
        self._code = 0
        self._skip_depth = 0
        self._heading: int | None = None
        self._list_stack: list[bool] = []  # True = ordenada
        self._pre_text: list[str] | None = None
        self._quote_depth = 0
        self._table_rows: list[list[list[_Run]]] | None = None
        self._cell_runs: list[_Run] | None = None

    # — helpers —
    def _open_paragraph(self) -> None:
        self._flush_paragraph()
        self._runs = []

    def _flush_paragraph(self) -> None:
        if self._runs is None:
            return
        runs = [r for r in self._runs if r.text]
        if runs:
            if self._heading:
                self.blocks.append(("heading", self._heading, runs))
            elif self._list_stack:
                self.blocks.append(
                    ("item", self._list_stack[-1], len(self._list_stack), runs)
                )
            elif self._quote_depth:
                self.blocks.append(("quote", runs))
            else:
                self.blocks.append(("para", runs))
        self._runs = None

    # — parser callbacks —
    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attr_map = {name: (value or "") for name, value in attrs}
        classes = attr_map.get("class", "")
        if self._skip_depth:
            if tag not in self._VOID:
                self._skip_depth += 1
            return
        if any(c in classes for c in self._SKIPPED) or tag == "nav":
            self._skip_depth += 1
            return
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._heading = int(tag[1])
            self._open_paragraph()
        elif tag == "p":
            if "page-break" in classes:
                self._flush_paragraph()
                self.blocks.append(("pagebreak",))
            else:
                self._open_paragraph()
        elif tag == "div" and "page-break" in classes:
            self._flush_paragraph()
            self.blocks.append(("pagebreak",))
        elif tag in ("ul", "ol"):
            self._flush_paragraph()
            self._list_stack.append(tag == "ol")
        elif tag == "li":
            self._open_paragraph()
        elif tag == "pre":
            self._flush_paragraph()
            self._pre_text = []
        elif tag == "blockquote":
            self._flush_paragraph()
            self._quote_depth += 1
        elif tag == "table":
            self._flush_paragraph()
            self._table_rows = []
        elif tag == "tr" and self._table_rows is not None:
            self._table_rows.append([])
        elif tag in ("th", "td") and self._table_rows is not None:
            self._cell_runs = []
            if tag == "th":
                self._bold += 1
        elif tag in ("strong", "b"):
            self._bold += 1
        elif tag in ("em", "i"):
            self._italic += 1
        elif tag == "code" and self._pre_text is None:
            self._code += 1
        elif tag == "br":
            self.handle_data("\n")
        elif tag == "img":
            image = _decode_data_uri(attr_map.get("src", ""))
            if image is not None:
                self._flush_paragraph()
                self.blocks.append(("image", image))

    def handle_endtag(self, tag: str) -> None:
        if self._skip_depth:
            self._skip_depth -= 1
            return
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._flush_paragraph()
            self._heading = None
        elif tag in ("p", "li"):
            self._flush_paragraph()
        elif tag in ("ul", "ol"):
            self._flush_paragraph()
            if self._list_stack:
                self._list_stack.pop()
        elif tag == "pre":
            if self._pre_text is not None:
                self.blocks.append(("codeblock", "".join(self._pre_text).rstrip("\n")))
            self._pre_text = None
        elif tag == "blockquote":
            self._flush_paragraph()
            self._quote_depth = max(0, self._quote_depth - 1)
        elif tag == "table":
            if self._table_rows:
                self.blocks.append(("table", self._table_rows))
            self._table_rows = None
        elif tag in ("th", "td"):
            if self._table_rows is not None and self._table_rows and self._cell_runs is not None:
                self._table_rows[-1].append(self._cell_runs)
            self._cell_runs = None
            if tag == "th":
                self._bold = max(0, self._bold - 1)
        elif tag in ("strong", "b"):
            self._bold = max(0, self._bold - 1)
        elif tag in ("em", "i"):
            self._italic = max(0, self._italic - 1)
        elif tag == "code" and self._pre_text is None:
            self._code = max(0, self._code - 1)

    def handle_data(self, data: str) -> None:
        if self._skip_depth:
            return
        if self._pre_text is not None:
            self._pre_text.append(data)
            return
        run = _Run(data, bold=self._bold > 0, italic=self._italic > 0, code=self._code > 0)
        if self._cell_runs is not None:
            self._cell_runs.append(run)
        else:
            if self._runs is None:
                if not data.strip():
                    return
                self._runs = []
            self._runs.append(run)

    def close(self) -> None:  # flush final
        super().close()
        self._flush_paragraph()


def _decode_data_uri(src: str) -> bytes | None:
    found = _DATA_URI.match(src.strip())
    if not found:
        return None
    try:
        return base64.b64decode(found.group("b64"), validate=False)
    except Exception:
        return None


def _component_text(html_body: str, component_class: str) -> str:
    """Texto plano del primer componente document-header/footer del documento."""
    for component in _RUNNING_COMPONENT.finditer(html_body):
        if component_class in component.group(0):
            return " ".join(_TAG_STRIP.sub(" ", component.group(0)).split())
    return ""


def _add_page_number_fields(paragraph: Any) -> None:
    """Añade "Página X de Y" con campos PAGE/NUMPAGES nativos de Word."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def field(instruction: str) -> list[Any]:
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        return [begin, instr, end]

    paragraph.add_run("Página ")
    for element in field("PAGE"):
        paragraph.add_run()._r.append(element)
    paragraph.add_run(" de ")
    for element in field("NUMPAGES"):
        paragraph.add_run()._r.append(element)


def _emit_runs(paragraph: Any, runs: list[_Run]) -> None:
    for r in runs:
        run = paragraph.add_run(r.text)
        run.bold = r.bold or None
        run.italic = r.italic or None
        if r.code:
            run.font.name = "Courier New"


def _list_style(document: Any, ordered: bool, depth: int) -> str:
    base = "List Number" if ordered else "List Bullet"
    name = base if depth <= 1 else f"{base} {min(depth, 3)}"
    try:
        document.styles[name]
        return name
    except KeyError:
        return base


def render_markdown_docx(
    markdown_text: str,
    variables: Mapping[str, str] | None = None,
    options: ExportOptions | None = None,
) -> bytes:
    """Markdown → DOCX (python-docx). Ver spec en el docstring del módulo."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.shared import Inches, Pt

    opts = options or ExportOptions()
    body, toc_tokens = _markdown_to_body(markdown_text)
    body = _resolve_variables(body, variables)

    document = Document()

    # Cabecera de sección: options manda; si no, el texto del componente del documento.
    header_text = opts.header_text or _component_text(body, "document-header")
    logo = _decode_data_uri(opts.logo_data_uri or "")
    if header_text or logo:
        header_paragraph = document.sections[0].header.paragraphs[0]
        if logo:
            try:
                header_paragraph.add_run().add_picture(io.BytesIO(logo), height=Pt(24))
                header_paragraph.add_run("  ")
            except Exception:
                pass  # logo corrupto: la cabecera de texto sigue valiendo
        header_paragraph.add_run(header_text)

    footer_paragraph = document.sections[0].footer.paragraphs[0]
    footer_text = opts.footer_text or _component_text(body, "document-footer")
    if footer_text:
        footer_paragraph.add_run(f"{footer_text}   ")
    _add_page_number_fields(footer_paragraph)
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Índice estático (sin números de página — ver docstring).
    if opts.include_toc:
        flat = _flatten_toc(toc_tokens)
        if flat:
            document.add_heading("Índice", level=1)
            for level, name in flat:
                toc_paragraph = document.add_paragraph(name)
                toc_paragraph.paragraph_format.left_indent = Inches(0.3 * (level - 2))
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    collector = _BlockCollector()
    collector.feed(body)
    collector.close()

    for block in collector.blocks:
        kind = block[0]
        if kind == "heading":
            _, level, runs = block
            _emit_runs(document.add_heading("", level=min(level, 4)), runs)
        elif kind == "para":
            _emit_runs(document.add_paragraph(), block[1])
        elif kind == "item":
            _, ordered, depth, runs = block
            _emit_runs(
                document.add_paragraph(style=_list_style(document, ordered, depth)), runs
            )
        elif kind == "quote":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.3)
            for r in block[1]:
                r.italic = True
            _emit_runs(paragraph, block[1])
        elif kind == "codeblock":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(block[1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif kind == "table":
            rows = block[1]
            cols = max((len(r) for r in rows), default=0)
            if cols == 0:
                continue
            table = document.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, cell_runs in enumerate(row):
                    _emit_runs(table.cell(i, j).paragraphs[0], cell_runs)
        elif kind == "image":
            try:
                document.add_picture(io.BytesIO(block[1]), width=Inches(5.9))
            except Exception:
                continue  # formato no soportado por python-docx: se omite la imagen
        elif kind == "pagebreak":
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
