"""
Servicio de plantillas/memorias de referencia (CompanyTemplate).

Responsable de:
  - Subir el archivo (PDF/DOCX) a Azure Blob Storage en un contenedor lógico
    aislado por usuario.
  - Extraer el texto del documento (Azure Document Intelligence para PDF,
    parser local para DOCX).
  - Generar una síntesis profunda con el agente de resumen
    (`COMPANY_TEMPLATE_SUMMARY_PROMPT`) que captura la estructura, tono, voz y
    propuesta de valor del documento original.
  - Persistir y exponer CRUD sobre el catálogo de plantillas por usuario.

El texto extraído (`extracted_text`) se guarda íntegro para poder regenerar el
resumen o auditar. El resumen (`summary`) es lo que se inyecta como contexto al
agente de Memoria Técnica — es la forma compacta de “transferir el alma” del
documento sin saturar el contexto del LLM.
"""

from __future__ import annotations

import io
import uuid
from datetime import datetime, timezone
from typing import Iterable

from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import AnalyzeResult
from azure.core.credentials import AzureKeyCredential
from azure.storage.blob import BlobServiceClient, ContentSettings
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.core.config import settings
from app.core.logging import get_logger
from app.models.domain import CompanyTemplate
from app.prompts.templates import COMPANY_TEMPLATE_SUMMARY_PROMPT
from app.services.embeddings import get_openai_client
from app.services.ingestion import validate_pdf_bytes

logger = get_logger(__name__)


# ── Configuración del agente de resumen ────────────────────────────────────────

SUMMARY_LLM_MODEL = "extraccion_datos_4o"
SUMMARY_TEMPERATURE = 0.3        # baja: queremos descripción fiel, no creatividad
# Tope de caracteres del texto extraído que mandamos al LLM para el resumen.
# Una memoria típica son 20–60 págs ≈ 30–90k chars. Cortamos a algo seguro para
# gpt-4o-mini (ventana ~128k tokens) reservando margen para la respuesta.
MAX_SUMMARY_INPUT_CHARS = 80_000
TEMPLATES_BLOB_PREFIX = "company-templates"  # subcarpeta lógica en el container

ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
}


class TemplateProcessingError(Exception):
    """Error de extracción/procesado de una plantilla."""


# ── Lectura ────────────────────────────────────────────────────────────────────

def list_templates(user_id: str, db: Session) -> list[CompanyTemplate]:
    """Plantillas del usuario, más recientes primero."""
    return (
        db.query(CompanyTemplate)
        .filter(CompanyTemplate.user_id == user_id)
        .order_by(CompanyTemplate.created_at.desc())
        .all()
    )


def get_template(template_id: str, user_id: str, db: Session) -> CompanyTemplate | None:
    return (
        db.query(CompanyTemplate)
        .filter(
            CompanyTemplate.id == template_id,
            CompanyTemplate.user_id == user_id,
        )
        .first()
    )


def get_templates_by_ids(
    template_ids: Iterable[str], user_id: str, db: Session
) -> list[CompanyTemplate]:
    """Recupera varias plantillas garantizando que pertenecen al user_id (§10)."""
    ids = [t for t in template_ids if t]
    if not ids:
        return []
    return (
        db.query(CompanyTemplate)
        .filter(
            CompanyTemplate.user_id == user_id,
            CompanyTemplate.id.in_(ids),
        )
        .all()
    )


# ── Subida a Blob Storage ──────────────────────────────────────────────────────

def _upload_to_blob(
    file_bytes: bytes, user_id: str, filename: str, mime_type: str
) -> str:
    """
    Sube el archivo al contenedor de pliegos bajo una subcarpeta lógica por usuario
    y devuelve la URL del blob (sin SAS). Reutilizamos el contenedor existente para
    no añadir infra nueva; el aislamiento se asegura por prefijo + filtro por user_id.
    """
    if not settings.AZURE_STORAGE_CONNECTION_STRING:
        raise TemplateProcessingError("Azure Storage no está configurado.")

    blob_service = BlobServiceClient.from_connection_string(
        settings.AZURE_STORAGE_CONNECTION_STRING
    )
    # nombre único para evitar colisiones; conservamos la extensión original
    safe_name = filename.replace("\\", "/").split("/")[-1]
    blob_name = f"{TEMPLATES_BLOB_PREFIX}/{user_id}/{uuid.uuid4()}-{safe_name}"

    blob_client = blob_service.get_blob_client(
        container=settings.AZURE_STORAGE_CONTAINER_NAME,
        blob=blob_name,
    )
    blob_client.upload_blob(
        file_bytes,
        overwrite=True,
        content_settings=ContentSettings(content_type=mime_type),
    )
    return blob_client.url


def _delete_blob(blob_url: str) -> None:
    if not settings.AZURE_STORAGE_CONNECTION_STRING or not blob_url:
        return
    try:
        blob_service = BlobServiceClient.from_connection_string(
            settings.AZURE_STORAGE_CONNECTION_STRING
        )
        container_part = f"/{settings.AZURE_STORAGE_CONTAINER_NAME}/"
        if container_part not in blob_url:
            return
        blob_name = blob_url.split(container_part, 1)[1]
        # los nombres en la URL vienen URL-encoded; decodificamos para el SDK
        from urllib.parse import unquote
        blob_name = unquote(blob_name)
        blob_service.get_blob_client(
            container=settings.AZURE_STORAGE_CONTAINER_NAME,
            blob=blob_name,
        ).delete_blob()
    except Exception as e:
        # No re-lanzamos: la limpieza del blob es best-effort tras borrar la DB.
        logger.warning(f"Could not delete template blob {blob_url}: {e}")


# ── Extracción de texto (PDF + DOCX) ───────────────────────────────────────────

def _extract_text_from_pdf(pdf_bytes: bytes, filename: str) -> tuple[str, int | None]:
    """Devuelve (texto, page_count). Usa Azure DI si está configurado; pypdf si no."""
    validate_pdf_bytes(pdf_bytes, filename)

    if settings.AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT and settings.AZURE_DOCUMENT_INTELLIGENCE_KEY:
        client = DocumentIntelligenceClient(
            endpoint=settings.AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT,
            credential=AzureKeyCredential(settings.AZURE_DOCUMENT_INTELLIGENCE_KEY),
        )
        poller = client.begin_analyze_document(
            "prebuilt-layout",
            body=io.BytesIO(pdf_bytes),
            content_type="application/pdf",
        )
        result: AnalyzeResult = poller.result()

        # Agrupamos por página para preservar saltos lógicos en el texto.
        page_texts: dict[int, list[str]] = {}
        for paragraph in (result.paragraphs or []):
            page_num = (
                paragraph.bounding_regions[0].page_number
                if paragraph.bounding_regions else 1
            )
            if paragraph.content.strip():
                page_texts.setdefault(page_num, []).append(paragraph.content.strip())
        for table in (result.tables or []):
            for cell in (table.cells or []):
                if cell.content and cell.content.strip():
                    page_num = (
                        cell.bounding_regions[0].page_number
                        if cell.bounding_regions else 1
                    )
                    page_texts.setdefault(page_num, []).append(cell.content.strip())

        text = "\n\n".join(
            "\n".join(page_texts[p]) for p in sorted(page_texts)
        )
        page_count = len(result.pages) if result.pages else None
        return text, page_count

    # Fallback local (sin DI configurado): pypdf básico.
    reader = PdfReader(io.BytesIO(pdf_bytes))
    pages = [(p.extract_text() or "").strip() for p in reader.pages]
    return "\n\n".join(t for t in pages if t), len(reader.pages)


def _extract_text_from_docx(docx_bytes: bytes) -> tuple[str, int | None]:
    """Extrae texto plano de un DOCX (párrafos + celdas de tabla). page_count no aplica."""
    try:
        from docx import Document  # python-docx
    except ImportError as e:
        raise TemplateProcessingError(
            "El soporte de DOCX requiere la dependencia `python-docx`. "
            "Añádela al backend antes de subir Word."
        ) from e

    doc = Document(io.BytesIO(docx_bytes))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if (c.text or "").strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts), None


def _extract_text(file_bytes: bytes, filename: str, mime_type: str) -> tuple[str, int | None]:
    if mime_type == "application/pdf":
        return _extract_text_from_pdf(file_bytes, filename)
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_text_from_docx(file_bytes)
    raise TemplateProcessingError(f"Tipo de archivo no soportado: {mime_type}")


# ── Agente de resumen profundo ─────────────────────────────────────────────────

async def _generate_summary(extracted_text: str, filename: str) -> str | None:
    """
    Llama al agente de síntesis profunda con el texto extraído.
    Devuelve el Markdown de la síntesis o None si el LLM no está configurado
    (la plantilla queda guardada con summary=None y reintentable luego).
    """
    if not extracted_text.strip():
        logger.warning(f"Empty extracted text for template '{filename}'; skipping summary.")
        return None

    client = get_openai_client()
    if not client:
        logger.warning("Azure OpenAI not configured — template summary unavailable.")
        return None

    # Truncamos por seguridad. El prompt instruye al modelo a tratar el texto como “documento”.
    text_input = extracted_text[:MAX_SUMMARY_INPUT_CHARS]
    if len(extracted_text) > MAX_SUMMARY_INPUT_CHARS:
        text_input += "\n\n[...documento truncado por longitud...]"

    user_message = (
        f"Documento de referencia: '{filename}'\n\n"
        f"TEXTO ÍNTEGRO DEL DOCUMENTO:\n{text_input}"
    )

    try:
        response = await client.chat.completions.create(
            model=SUMMARY_LLM_MODEL,
            temperature=SUMMARY_TEMPERATURE,
            messages=[
                {"role": "system", "content": COMPANY_TEMPLATE_SUMMARY_PROMPT},
                {"role": "user", "content": user_message},
            ],
        )
        summary = (response.choices[0].message.content or "").strip()
        return summary or None
    except Exception as e:
        logger.error(
            f"Error generating template summary for '{filename}': {e}",
            exc_info=True,
        )
        return None


# ── Flujo principal: subida + extracción + resumen + persistencia ──────────────

async def create_template_from_upload(
    user_id: str,
    file_bytes: bytes,
    filename: str,
    mime_type: str,
    title: str | None,
    description: str | None,
    db: Session,
) -> CompanyTemplate:
    """
    Procesa una plantilla nueva subida por el usuario:
      1. Sube el archivo a Blob Storage.
      2. Extrae el texto (DI para PDF, python-docx para DOCX).
      3. Llama al agente de resumen profundo (no bloquea la creación si falla).
      4. Persiste el registro en `company_templates`.

    Devuelve la plantilla ya persistida (con `summary` poblado si el LLM respondió).
    """
    if mime_type not in ALLOWED_MIME_TYPES:
        raise TemplateProcessingError(
            f"Tipo de archivo no soportado: {mime_type}. Solo PDF o DOCX."
        )
    if not file_bytes:
        raise TemplateProcessingError("El archivo está vacío.")

    file_size = len(file_bytes)

    # 1) Blob Storage
    blob_url = _upload_to_blob(file_bytes, user_id, filename, mime_type)

    # 2) Extracción de texto
    try:
        extracted_text, page_count = _extract_text(file_bytes, filename, mime_type)
    except Exception as e:
        _delete_blob(blob_url)
        logger.error(f"Failed to extract text from template '{filename}': {e}", exc_info=True)
        raise TemplateProcessingError(f"No se pudo extraer texto del documento: {e}")

    if not extracted_text.strip():
        _delete_blob(blob_url)
        raise TemplateProcessingError(
            "No se pudo extraer texto del documento. Revisa que no esté vacío ni protegido."
        )

    # 3) Resumen profundo (best-effort; si falla, persistimos sin summary)
    summary = await _generate_summary(extracted_text, filename)

    # 4) Persistencia
    now = datetime.now(timezone.utc)
    template = CompanyTemplate(
        id=str(uuid.uuid4()),
        user_id=user_id,
        filename=filename,
        title=(title or "").strip() or None,
        description=(description or "").strip() or None,
        mime_type=mime_type,
        file_size=file_size,
        page_count=page_count,
        blob_url=blob_url,
        extracted_text=extracted_text,
        summary=summary,
        created_at=now,
    )
    db.add(template)
    db.commit()
    db.refresh(template)

    logger.info(
        "Company template created",
        extra={
            "template_id": template.id,
            "user_id": user_id,
            "template_filename": filename,
            "template_mime_type": mime_type,
            "template_file_size": file_size,
            "template_page_count": page_count,
            "summary_generated": summary is not None,
            "summary_len": len(summary) if summary else 0,
            "extracted_len": len(extracted_text),
        },
    )
    return template


def delete_template(template_id: str, user_id: str, db: Session) -> bool:
    """Borra una plantilla (DB + blob). True si existía y se borró."""
    template = get_template(template_id, user_id, db)
    if template is None:
        return False
    blob_url = template.blob_url
    db.delete(template)
    db.commit()
    _delete_blob(blob_url)
    logger.info(
        "Company template deleted",
        extra={"template_id": template_id, "user_id": user_id},
    )
    return True


def update_template_metadata(
    template_id: str,
    user_id: str,
    title: str | None,
    description: str | None,
    db: Session,
) -> CompanyTemplate | None:
    """Edita solo metadatos visibles (no toca extracted_text/summary/blob)."""
    template = get_template(template_id, user_id, db)
    if template is None:
        return None
    if title is not None:
        template.title = title.strip() or None
    if description is not None:
        template.description = description.strip() or None
    db.commit()
    db.refresh(template)
    return template


# ── Inyección en el prompt del agente de Memoria Técnica ───────────────────────

TEMPLATE_CONTEXT_MAX_CHARS = 18_000  # límite por bloque inyectado, no por plantilla


def build_templates_context(templates: list[CompanyTemplate]) -> str:
    """
    Construye el bloque `<PlantillasDeReferencia>` que se inyecta en los prompts
    de esquema y propuesta. Prefiere el `summary` (síntesis profunda); si una
    plantilla no tiene resumen aún, cae al texto extraído truncado.
    """
    if not templates:
        return ""

    parts: list[str] = ["<PlantillasDeReferencia>"]
    for i, t in enumerate(templates, start=1):
        label = t.title or t.filename
        header = f"\n### Plantilla {i}: {label}"
        if t.description:
            header += f"\nDescripción del usuario: {t.description}"

        body = (t.summary or "").strip()
        if body:
            body = f"\n\n{body}"
        else:
            # Fallback si aún no hay resumen: texto truncado (mejor algo que nada).
            raw = (t.extracted_text or "").strip()
            if not raw:
                continue
            truncated = raw[:TEMPLATE_CONTEXT_MAX_CHARS]
            if len(raw) > TEMPLATE_CONTEXT_MAX_CHARS:
                truncated += "\n[...truncado...]"
            body = f"\n\n[Resumen no disponible — texto bruto del documento:]\n{truncated}"

        parts.append(header + body)
    parts.append("\n</PlantillasDeReferencia>")
    return "\n".join(parts)
